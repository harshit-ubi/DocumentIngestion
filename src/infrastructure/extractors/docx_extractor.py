import io
import docx
from src.domain.interfaces.extractor import DocumentExtractorInterface
from src.core.exceptions import ExtractionError
from src.core.logging import logger


class DocxExtractor(DocumentExtractorInterface):
    """DOCX Document Extractor using python-docx."""

    def extract_text(self, file_content: bytes, filename: str) -> str:
        try:
            docx_file = io.BytesIO(file_content)
            doc = docx.Document(docx_file)
            extracted_paragraphs = []

            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    extracted_paragraphs.append(p.text.strip())

            # Also extract text inside tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        extracted_paragraphs.append(row_text)

            if not extracted_paragraphs:
                raise ExtractionError(f"No extractable text content found in DOCX '{filename}'.")

            logger.info(f"Successfully extracted {len(extracted_paragraphs)} paragraphs/rows from DOCX '{filename}'.")
            return "\n\n".join(extracted_paragraphs)
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX '{filename}': {str(e)}")
            raise ExtractionError(f"Failed to extract text from DOCX '{filename}': {str(e)}")
